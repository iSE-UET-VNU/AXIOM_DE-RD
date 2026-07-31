from __future__ import annotations

import importlib.util
import json
import subprocess
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from src.artifacts import write_ingested_artifacts
from src.ingestion.parsing.backends import WordConfig, WordParserBackend
from src.ingestion.parsing.chandra2 import Chandra2Provider
from src.ingestion.runner import run as run_ingestion
from src.models import DataObject, ParseStatus, PipelineState


DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
PIL_AVAILABLE = importlib.util.find_spec("PIL") is not None


def _data_object(path: Path, object_id: str = "word-document") -> DataObject:
    return DataObject(
        object_id=object_id,
        uri=path.as_posix(),
        content_type="application/msword" if path.suffix == ".doc" else (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        metadata={"format": path.suffix.lstrip(".").lower(), "file_name": path.name},
    )


def _add_hyperlink(paragraph: object, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    relationship_id = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


@unittest.skipUnless(DOCX_AVAILABLE and PIL_AVAILABLE, "Word test dependencies missing")
class WordParserTests(unittest.TestCase):
    def _document(self, root: Path) -> Path:
        from docx import Document
        from docx.shared import Inches
        from PIL import Image

        large_image = root / "large.png"
        small_image = root / "small.png"
        Image.new("RGB", (400, 300), color=(30, 120, 190)).save(large_image)
        Image.new("RGB", (20, 20), color=(190, 30, 30)).save(small_image)

        document = Document()
        document.core_properties.title = "Báo cáo thử nghiệm"
        document.add_heading("Tiêu đề Heading", level=1)
        document.add_paragraph("Nội dung Unicode tiếng Việt.")
        hyperlink_paragraph = document.add_paragraph("Tham khảo ")
        _add_hyperlink(hyperlink_paragraph, "cổng dữ liệu", "https://example.com")
        bullet = document.add_paragraph("Mục danh sách", style="List Bullet")
        bullet._p.get_or_add_pPr().get_or_add_numPr()

        table = document.add_table(rows=3, cols=3)
        values = [
            ["Tên gốc", "", "Điểm"],
            ["Alpha", "A2", "91%"],
            ["Beta", "B2", "88%"],
        ]
        for row_index, row in enumerate(values):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = value
        table.cell(0, 0).merge(table.cell(0, 1))

        picture_paragraph = document.add_paragraph()
        large = picture_paragraph.add_run().add_picture(
            str(large_image), width=Inches(4), height=Inches(3)
        )
        large._inline.docPr.set("descr", "Sơ đồ kiến trúc")
        document.add_paragraph().add_run().add_picture(
            str(large_image), width=Inches(4), height=Inches(3)
        )
        table.cell(2, 1).paragraphs[0].add_run().add_picture(
            str(small_image), width=Inches(0.1), height=Inches(0.1)
        )

        document.sections[0].header.paragraphs[0].text = "Không lấy header"
        path = root / "sample.docx"
        document.save(path)
        return path

    def test_extracts_continuous_text_tables_images_and_stable_blocks(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = self._document(root)
            parser = WordParserBackend(
                WordConfig(output_dir=str(root / "assets"), project_root=str(root))
            )
            first = parser.parse(path, _data_object(path))
            second = parser.parse(path, _data_object(path))

            self.assertEqual(first.status, ParseStatus.SUCCESS, first.error)
            parsed = first.parsed_data
            self.assertIsNotNone(parsed)
            row = parsed.rows[0]
            extraction = row["extraction"]
            self.assertEqual(parsed.text, row["text"])
            self.assertEqual(row["text"], extraction["main_text"])
            self.assertIn("Nội dung Unicode tiếng Việt.", parsed.text)
            self.assertIn("Tham khảo cổng dữ liệu", parsed.text)
            self.assertNotIn("Không lấy header", parsed.text)
            self.assertIn("\n\n", parsed.text)
            self.assertEqual(extraction["document_type"], "document")
            self.assertEqual(extraction["title"], "Báo cáo thử nghiệm")

            self.assertEqual(len(extraction["tables"]), 1)
            extracted_table = extraction["tables"][0]
            self.assertTrue(extracted_table["content"].startswith("| Tên gốc |  | Điểm |"))
            self.assertEqual(extracted_table["headers"], ["Tên gốc", "column_2", "Điểm"])
            self.assertEqual(extracted_table["rows"][0], ["Alpha", "A2", "91%"])
            self.assertEqual(parsed.tables[0].source_ref, extracted_table["source_ref"])

            self.assertEqual(len(extraction["figures"]), 3)
            self.assertIn(
                "Sơ đồ kiến trúc",
                [figure["description"] for figure in extraction["figures"]],
            )
            assets = parsed.metadata["image_files"]
            self.assertEqual(len(assets), 3)
            self.assertEqual(assets[1]["path"], assets[2]["path"])
            self.assertNotEqual(assets[0]["source_ref"], assets[1]["source_ref"])
            self.assertEqual(
                len({block["component_id"] for block in row["source_blocks"]}),
                len(row["source_blocks"]),
            )
            self.assertEqual(
                row["reading_order"],
                second.parsed_data.rows[0]["reading_order"],
            )
            self.assertIsNone(parsed.metadata["page_count"])
            self.assertEqual(parsed.metadata["pagination_source"], "unavailable")

    def test_ingestion_routes_locally_and_filters_small_images(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            path = self._document(root)
            assets = root / "ingested" / "assets"
            with patch.object(
                Chandra2Provider,
                "parse_file",
                side_effect=AssertionError("Chandra2 must not receive DOCX"),
            ):
                output = run_ingestion(
                    path,
                    parser_config={
                        "provider": "chandra2",
                        "word": {
                            "output_dir": str(assets),
                            "project_root": str(root),
                        },
                    },
                    project_root=root,
                )

            parsed = output.parsed_data[0]
            report = parsed.metadata["image_filtering"]
            self.assertEqual(parsed.metadata["backend"], "word")
            self.assertEqual(report["before_count"], 3)
            self.assertEqual(report["kept_count"], 2)
            self.assertEqual(report["dropped_count"], 1)
            self.assertEqual(len(parsed.rows[0]["extraction"]["figures"]), 2)
            self.assertEqual(len(parsed.metadata["image_files"]), 2)
            for asset in parsed.metadata["image_files"]:
                filtered_path = root / asset["path"]
                self.assertTrue(filtered_path.is_file())
                self.assertEqual(filtered_path.parent.name, "filtered_images")

            state = PipelineState(
                run_id="word-test",
                input_source=path.as_posix(),
                embedded_dir="embedded",
                output_dir="output",
                data_objects=output.data_objects,
                parsed_data=output.parsed_data,
                initial_schemas=output.initial_schemas,
                ingestion_config={"provider": "chandra2"},
            )
            artifact_dir = root / "ingested"
            write_ingested_artifacts(state, artifact_dir, project_root=root)
            document_path = next((artifact_dir / "documents").glob("*.json"))
            payload = json.loads(document_path.read_text(encoding="utf-8"))

            self.assertEqual(payload["contract_version"], "ingested-document-v2")
            artifact_extraction = payload["parsed"]["rows"][0]["extraction"]
            self.assertEqual(artifact_extraction["main_text"], parsed.text)
            self.assertEqual(len(artifact_extraction["tables"]), 1)
            self.assertEqual(len(artifact_extraction["figures"]), 2)

    def test_corrupt_docx_returns_failed_result(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "broken.docx"
            path.write_bytes(b"not an Open XML package")
            result = WordParserBackend().parse(path, _data_object(path))

        self.assertEqual(result.status, ParseStatus.FAILED)
        self.assertTrue(result.error["message"])

    def test_one_row_table_and_empty_document_are_valid(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            one_row_path = root / "one-row.docx"
            one_row = Document()
            one_row.core_properties.title = "Word Document"
            table = one_row.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Tên"
            table.cell(0, 1).text = "Tên"
            one_row.add_heading("Tiêu đề thật", level=1)
            one_row.save(one_row_path)

            empty_path = root / "empty.docx"
            empty = Document()
            empty.add_table(rows=1, cols=2)
            empty.save(empty_path)

            one_row_result = WordParserBackend().parse(
                one_row_path, _data_object(one_row_path)
            )
            empty_result = WordParserBackend().parse(
                empty_path, _data_object(empty_path)
            )

        self.assertEqual(one_row_result.status, ParseStatus.SUCCESS)
        extracted = one_row_result.parsed_data.rows[0]["extraction"]["tables"][0]
        self.assertEqual(extracted["headers"], ["Tên", "Tên_2"])
        self.assertEqual(extracted["rows"], [])
        self.assertTrue(extracted["content"].startswith("| Tên | Tên |"))
        self.assertEqual(
            one_row_result.parsed_data.rows[0]["extraction"]["title"],
            "Tiêu đề thật",
        )
        self.assertEqual(empty_result.status, ParseStatus.SUCCESS)
        self.assertEqual(empty_result.parsed_data.text, "")
        self.assertEqual(empty_result.parsed_data.tables, [])
        self.assertEqual(
            empty_result.parsed_data.rows[0]["extraction"]["tables"], []
        )
        self.assertIsNone(empty_result.parsed_data.rows[0]["extraction"]["title"])


@unittest.skipUnless(DOCX_AVAILABLE, "python-docx missing")
class LegacyDocConversionTests(unittest.TestCase):
    def test_successful_conversion_preserves_original_provenance(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "legacy.doc"
            source.write_bytes(b"\xd0\xcf\x11\xe0")

            def fake_run(command: list[str], **kwargs: object):
                profile_argument = next(
                    item for item in command if item.startswith("-env:UserInstallation=")
                )
                self.assertIn("/profile", profile_argument)
                self.assertIs(kwargs["shell"], False)
                output_dir = Path(command[command.index("--outdir") + 1])
                converted = Document()
                converted.add_paragraph("Legacy text")
                converted.save(output_dir / "legacy.docx")
                return subprocess.CompletedProcess(command, 0, "converted", "")

            with patch(
                "src.ingestion.parsing.backends.word._find_soffice",
                return_value=Path("soffice.com"),
            ), patch(
                "src.ingestion.parsing.backends.word.subprocess.run",
                side_effect=fake_run,
            ):
                result = WordParserBackend().parse(source, _data_object(source))

        self.assertEqual(result.status, ParseStatus.SUCCESS, result.error)
        parsed = result.parsed_data
        self.assertEqual(parsed.source_uri, source.as_posix())
        self.assertEqual(parsed.source_format, "doc")
        self.assertEqual(parsed.text, "Legacy text")
        self.assertEqual(parsed.metadata["converted_from"], "doc")
        self.assertEqual(parsed.metadata["conversion"]["converter"], "libreoffice")

    def test_missing_converter_timeout_exit_and_missing_output_are_actionable(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "legacy.doc"
            source.write_bytes(b"\xd0\xcf\x11\xe0")
            data_object = _data_object(source)

            with patch(
                "src.ingestion.parsing.backends.word._find_soffice",
                return_value=None,
            ):
                missing = WordParserBackend().parse(source, data_object)
            with patch(
                "src.ingestion.parsing.backends.word._find_soffice",
                return_value=Path("soffice.com"),
            ), patch(
                "src.ingestion.parsing.backends.word.subprocess.run",
                side_effect=subprocess.TimeoutExpired("soffice", 1),
            ):
                timeout = WordParserBackend(
                    WordConfig(conversion_timeout_seconds=1)
                ).parse(source, data_object)
            with patch(
                "src.ingestion.parsing.backends.word._find_soffice",
                return_value=Path("soffice.com"),
            ), patch(
                "src.ingestion.parsing.backends.word.subprocess.run",
                return_value=subprocess.CompletedProcess([], 7, "", "conversion failed"),
            ):
                failed_exit = WordParserBackend().parse(source, data_object)
            with patch(
                "src.ingestion.parsing.backends.word._find_soffice",
                return_value=Path("soffice.com"),
            ), patch(
                "src.ingestion.parsing.backends.word.subprocess.run",
                return_value=subprocess.CompletedProcess([], 0, "converted", ""),
            ):
                no_output = WordParserBackend().parse(source, data_object)

        for result in (missing, timeout, failed_exit, no_output):
            self.assertEqual(result.status, ParseStatus.FAILED)
            self.assertTrue(result.error["message"])
        self.assertIn("required", missing.error["message"])
        self.assertIn("timed out", timeout.error["message"])
        self.assertIn("exit 7", failed_exit.error["message"])
        self.assertIn("did not create", no_output.error["message"])


if __name__ == "__main__":
    unittest.main()
