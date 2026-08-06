"""Text extraction from lake files, one handler per format.

Extraction preserves enough structure for the structure-aware arms of the
benchmark: paragraphs keep their page, tables emit one block per row with the
header carried alongside. Arms that only need flat text join ``blocks``.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any, Callable
import csv as csvlib
import re

_HANDLERS: dict[str, Callable[[Path], list["Block"]]] = {}
MAX_TABLE_ROWS = 2000
MAX_CHARS = 2_000_000


@dataclass
class Block:
    text: str
    kind: str = "paragraph"
    page: int | None = None
    section: str | None = None


@dataclass
class ExtractedDoc:
    doc_id: str
    title: str
    blocks: list[Block] = field(default_factory=list)
    error: str | None = None
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks if block.text.strip())

    @property
    def ok(self) -> bool:
        return self.error is None and any(block.text.strip() for block in self.blocks)


def handler(*extensions: str) -> Callable[[Callable[[Path], list[Block]]], Callable[[Path], list[Block]]]:
    def register(fn: Callable[[Path], list[Block]]) -> Callable[[Path], list[Block]]:
        for extension in extensions:
            _HANDLERS[extension] = fn
        return fn

    return register


def supported_extensions() -> frozenset[str]:
    return frozenset(_HANDLERS)


def extract(path: Path | str, doc_id: str | None = None) -> ExtractedDoc:
    path = Path(path)
    identifier = doc_id or path.name
    title = _title_from_name(path.stem)
    fn = _HANDLERS.get(path.suffix.lower())
    if fn is None:
        return ExtractedDoc(identifier, title, error=f"unsupported:{path.suffix.lower()}")
    try:
        blocks = fn(path)
    except Exception as exc:
        return ExtractedDoc(identifier, title, error=f"{type(exc).__name__}: {exc}")
    return ExtractedDoc(identifier, title, blocks=_clip(blocks))


@handler(".epub")
def _epub(path: Path) -> list[Block]:
    """EPUB is a zip of XHTML, read in spine order.

    Parsed with the standard library rather than ebooklib: the format is a zip
    plus a manifest, adding a dependency buys nothing, and the spine is the part
    that matters. Reading files in archive order instead of spine order would
    interleave chapters with front matter and scramble a novel's reading order --
    the same failure the pipeline's reading-order work exists to prevent.
    """
    import zipfile
    from xml.etree import ElementTree

    blocks: list[Block] = []
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        opf_name = _epub_opf(archive, names)
        if opf_name is None:
            return blocks
        opf = ElementTree.fromstring(archive.read(opf_name))
        base = opf_name.rsplit("/", 1)[0] + "/" if "/" in opf_name else ""

        ns = {"opf": "http://www.idpf.org/2007/opf"}
        hrefs = {
            item.get("id"): item.get("href")
            for item in opf.iterfind(".//opf:manifest/opf:item", ns)
        }
        spine = [
            hrefs.get(ref.get("idref"))
            for ref in opf.iterfind(".//opf:spine/opf:itemref", ns)
        ]
        for href in spine:
            if not href:
                continue
            name = _epub_resolve(base + href, names)
            if name is None:
                continue
            parser = _EpubText()
            parser.feed(archive.read(name).decode("utf-8", errors="replace"))
            blocks.extend(parser.blocks)
    return blocks


def _epub_opf(archive: Any, names: set[str]) -> str | None:
    """Locate the package document via META-INF/container.xml, with a fallback."""
    from xml.etree import ElementTree

    if "META-INF/container.xml" in names:
        container = ElementTree.fromstring(archive.read("META-INF/container.xml"))
        for root in container.iter():
            if root.tag.endswith("rootfile") and root.get("full-path"):
                return root.get("full-path")
    return next((n for n in sorted(names) if n.endswith(".opf")), None)


def _epub_resolve(name: str, names: set[str]) -> str | None:
    """Normalize ``a/../b`` and drop fragments; hrefs are URL-relative."""
    name = name.split("#", 1)[0]
    parts: list[str] = []
    for part in name.split("/"):
        if part == "..":
            if parts:
                parts.pop()
        elif part not in ("", "."):
            parts.append(part)
    resolved = "/".join(parts)
    return resolved if resolved in names else None


class _EpubText(HTMLParser):
    """Headings and paragraphs out of one XHTML document."""

    _BLOCK = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "div", "li", "blockquote"}
    _SKIP = {"script", "style", "head"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[Block] = []
        self._buffer: list[str] = []
        self._kind = "paragraph"
        self._skipping = 0

    def handle_starttag(self, tag: str, attrs: Any) -> None:
        if tag in self._SKIP:
            self._skipping += 1
        elif tag in self._BLOCK:
            self._flush()
            self._kind = "heading" if tag.startswith("h") and tag != "hr" else "paragraph"

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP:
            self._skipping = max(0, self._skipping - 1)
        elif tag in self._BLOCK:
            self._flush()

    def handle_data(self, data: str) -> None:
        if not self._skipping and data.strip():
            self._buffer.append(data)

    def _flush(self) -> None:
        text = " ".join(" ".join(self._buffer).split())
        self._buffer.clear()
        if text:
            self.blocks.append(Block(text, self._kind))
        self._kind = "paragraph"

    def close(self) -> None:  # noqa: D102
        super().close()
        self._flush()


@handler(".txt", ".md")
def _plain(path: Path) -> list[Block]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return [Block(part.strip(), _kind_of(part)) for part in _split_paragraphs(text)]


@handler(".pdf")
def _pdf(path: Path) -> list[Block]:
    import fitz

    blocks: list[Block] = []
    with fitz.open(path) as document:
        for number, page in enumerate(document, start=1):
            for part in _split_paragraphs(page.get_text("text") or ""):
                blocks.append(Block(part.strip(), _kind_of(part), page=number))
    return blocks


@handler(".docx")
def _docx(path: Path) -> list[Block]:
    import docx

    document = docx.Document(str(path))
    blocks: list[Block] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        blocks.append(Block(text, "heading" if "heading" in style or "title" in style else "paragraph"))
    for table in document.tables:
        blocks.extend(_table_blocks([[cell.text for cell in row.cells] for row in table.rows]))
    return blocks


@handler(".csv")
def _csv(path: Path) -> list[Block]:
    with path.open(encoding="utf-8-sig", errors="replace", newline="") as handle:
        sample = handle.read(8192)
        handle.seek(0)
        try:
            dialect: Any = csvlib.Sniffer().sniff(sample) if sample.strip() else csvlib.excel
        except csvlib.Error:
            dialect = csvlib.excel
        rows = [row for _, row in zip(range(MAX_TABLE_ROWS + 1), csvlib.reader(handle, dialect))]
    return _table_blocks(rows)


@handler(".xlsx", ".xlsm")
def _xlsx(path: Path) -> list[Block]:
    import openpyxl

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    blocks: list[Block] = []
    try:
        for sheet in workbook.worksheets:
            rows: list[list[Any]] = []
            for index, row in enumerate(sheet.iter_rows(values_only=True)):
                if index > MAX_TABLE_ROWS:
                    break
                rows.append(list(row))
            if not rows:
                continue
            blocks.append(Block(f"Sheet: {sheet.title}", "heading", section=sheet.title))
            for block in _table_blocks(rows):
                block.section = sheet.title
                blocks.append(block)
    finally:
        workbook.close()
    return blocks


HEADER_SCAN_ROWS = 12


def _is_numeric(value: str) -> bool:
    return bool(value) and value.replace(",", "").replace(".", "").replace("-", "").replace("%", "").isdigit()


def _header_score(row: list[str], below: list[list[str]]) -> float:
    filled = [cell for cell in row if cell]
    if len(filled) < 2:
        return -1.0
    labels = sum(1 for cell in filled if not _is_numeric(cell) and len(cell) <= 60)
    unique = len({cell.lower() for cell in filled}) / len(filled)
    width = len(filled) / max(len(row), 1)
    # A header names columns that the rows beneath actually fill with values.
    support = 0.0
    if below:
        widths = [sum(1 for cell in r if cell) for r in below[:20]]
        support = min(1.0, (sum(widths) / len(widths)) / max(len(filled), 1))
    numeric_below = 0.0
    if below:
        cells = [cell for r in below[:20] for cell in r if cell]
        numeric_below = sum(1 for cell in cells if _is_numeric(cell)) / max(len(cells), 1)
    return (labels / len(filled)) + unique + width + support + numeric_below


def _find_header(cleaned: list[list[str]]) -> int:
    best, best_score = 0, -2.0
    for index in range(min(HEADER_SCAN_ROWS, len(cleaned))):
        score = _header_score(cleaned[index], cleaned[index + 1 :])
        if score > best_score:
            best, best_score = index, score
    return best


def _table_blocks(rows: list[list[Any]]) -> list[Block]:
    cleaned = [[_cell(value) for value in row] for row in rows if any(_cell(v) for v in row)]
    if not cleaned:
        return []
    start = _find_header(cleaned)
    blocks = [Block(" ".join(c for c in row if c), "paragraph") for row in cleaned[:start]]
    header = cleaned[start]
    blocks.append(Block(" | ".join(cell for cell in header if cell), "table_header"))
    for row in cleaned[start + 1 :]:
        pairs = [f"{name}: {value}" for name, value in zip(header, row) if value and name]
        blocks.append(Block("; ".join(pairs) if pairs else " | ".join(c for c in row if c), "table_row"))
    return blocks


def _cell(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _split_paragraphs(text: str) -> list[str]:
    return [part for part in re.split(r"\n\s*\n", text or "") if part.strip()]


def _kind_of(part: str) -> str:
    stripped = part.strip()
    if stripped.startswith("#"):
        return "heading"
    if len(stripped) < 80 and stripped == stripped.upper() and any(c.isalpha() for c in stripped):
        return "heading"
    return "paragraph"


def _title_from_name(stem: str) -> str:
    return re.sub(r"[_\-.]+", " ", stem).strip()


_SEPARATORS = re.compile("[\u2028\u2029\x0b\x0c\x1c\x1d\x1e\x85]+")


def _strip_separators(text: str) -> str:
    return _SEPARATORS.sub(" ", text)


def _clip(blocks: list[Block]) -> list[Block]:
    total = 0
    kept: list[Block] = []
    for block in blocks:
        # U+2028/U+2029 survive json.dumps but break str.splitlines() on read.
        text = re.sub(r"[ \t]+", " ", _strip_separators(block.text or "")).strip()
        if not text:
            continue
        block.text = text
        total += len(text)
        kept.append(block)
        if total >= MAX_CHARS:
            break
    return kept
