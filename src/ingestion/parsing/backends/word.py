"""Native Word parser with optional LibreOffice conversion for legacy DOC."""

from __future__ import annotations

from dataclasses import dataclass
import hashlib
import mimetypes
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import time
from typing import Any, Iterable

from ....models import DataObject, ParsedData, ParsedTable, ParseResult
from ....utils.paths import portable_path
from ..contracts import AXIOM_NATIVE_BLOCK_SOURCE
from ._tabular import normalize_headers


WORD_EXTENSIONS = frozenset({".doc", ".docx"})
_CONTENT_ELEMENTS = frozenset({"object", "pict", "txbxContent", "altChunk"})
_GENERIC_CORE_TITLES = frozenset({"document", "word document"})


@dataclass(frozen=True)
class WordConfig:
    """Settings for native DOCX parsing and legacy DOC conversion."""

    extract_images: bool = True
    soffice_path: str | None = None
    conversion_timeout_seconds: float = 120.0
    output_dir: str | None = None
    project_root: str | None = None

    @classmethod
    def from_mapping(cls, config: dict[str, Any] | None) -> "WordConfig":
        config = config or {}
        timeout = float(config.get("conversion_timeout_seconds", 120))
        if timeout <= 0:
            raise ValueError("parsing.word.conversion_timeout_seconds must be positive")
        return cls(
            extract_images=bool(config.get("extract_images", True)),
            soffice_path=_optional_str(config.get("soffice_path")),
            conversion_timeout_seconds=timeout,
            output_dir=_optional_str(config.get("output_dir")),
            project_root=_optional_str(config.get("project_root")),
        )


@dataclass
class _WordState:
    source_blocks: list[dict[str, Any]]
    text_parts: list[str]
    text_citations: list[str]
    tables: list[dict[str, Any]]
    parsed_tables: list[ParsedTable]
    figures: list[dict[str, Any]]
    image_files: list[dict[str, Any]]
    unsupported_elements: dict[str, int]
    heading_title: tuple[str, str] | None = None
    first_paragraph: tuple[str, str] | None = None


class WordParserBackend:
    """Parse DOCX locally and DOC after a temporary LibreOffice conversion."""

    supported_extensions = WORD_EXTENSIONS
    backend_name = "word"

    def __init__(self, config: WordConfig | None = None) -> None:
        self.config = config or WordConfig()

    def parse(self, path: str | Path, data_object: DataObject) -> ParseResult:
        source = Path(path)
        try:
            if source.suffix.lower() == ".doc":
                parsed = self._parse_doc(source, data_object)
            else:
                parsed = self._parse_docx(source, data_object, source_format="docx")
        except Exception as exc:
            return ParseResult.failed(
                data_object.object_id,
                self.backend_name,
                exc,
                route=self.backend_name,
            )
        return ParseResult.success(
            data_object.object_id,
            self.backend_name,
            parsed,
            route=self.backend_name,
        )

    def _parse_doc(self, path: Path, data_object: DataObject) -> ParsedData:
        executable = _find_soffice(self.config.soffice_path)
        if executable is None:
            raise RuntimeError(
                "LibreOffice is required to parse legacy .doc files. Install "
                "LibreOffice, add soffice to PATH, or set parsing.word.soffice_path."
            )

        started = time.monotonic()
        with tempfile.TemporaryDirectory(prefix="axiom-word-") as temp_name:
            temp_dir = Path(temp_name)
            output_dir = temp_dir / "output"
            profile_dir = temp_dir / "profile"
            output_dir.mkdir()
            profile_dir.mkdir()
            command = [
                str(executable),
                f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
                "--headless",
                "--nologo",
                "--nodefault",
                "--nofirststartwizard",
                "--convert-to",
                "docx:Office Open XML Text",
                "--outdir",
                str(output_dir),
                str(path.resolve()),
            ]
            try:
                completed = subprocess.run(
                    command,
                    capture_output=True,
                    text=True,
                    errors="replace",
                    timeout=self.config.conversion_timeout_seconds,
                    check=False,
                    shell=False,
                )
            except subprocess.TimeoutExpired as exc:
                raise RuntimeError(
                    "LibreOffice conversion timed out after "
                    f"{self.config.conversion_timeout_seconds:g} seconds for {path.name}."
                ) from exc
            if completed.returncode != 0:
                detail = (completed.stderr or completed.stdout or "").strip()
                raise RuntimeError(
                    f"LibreOffice failed to convert {path.name} "
                    f"(exit {completed.returncode}): {detail or 'no diagnostic output'}"
                )

            converted = _converted_docx(output_dir, path.stem)
            if converted is None:
                detail = (completed.stderr or completed.stdout or "").strip()
                raise RuntimeError(
                    f"LibreOffice reported success but did not create a DOCX for "
                    f"{path.name}: {detail or 'no diagnostic output'}"
                )
            parsed = self._parse_docx(converted, data_object, source_format="doc")

        parsed.metadata["converted_from"] = "doc"
        parsed.metadata["conversion"] = {
            "converter": "libreoffice",
            "executable": str(executable),
            "elapsed_seconds": round(time.monotonic() - started, 3),
        }
        return parsed

    def _parse_docx(
        self,
        path: Path,
        data_object: DataObject,
        *,
        source_format: str,
    ) -> ParsedData:
        try:
            from docx import Document
            from docx.table import Table
        except ImportError as exc:
            raise RuntimeError(
                "python-docx>=1.2,<2 is required to parse Word files"
            ) from exc

        try:
            document = Document(str(path))
        except Exception as exc:
            raise RuntimeError(
                f"Unable to open {path.name} as DOCX; the file may be corrupt "
                "or encrypted."
            ) from exc
        state = _WordState([], [], [], [], [], [], [], {})
        saved_assets: dict[str, dict[str, Any]] = {}
        page_box = _page_box(document)

        for item in document.iter_inner_content():
            if isinstance(item, Table):
                self._parse_table(
                    item,
                    document=document,
                    page_box=page_box,
                    data_object=data_object,
                    state=state,
                    saved_assets=saved_assets,
                )
            else:
                self._parse_paragraph(
                    item,
                    document=document,
                    page_box=page_box,
                    data_object=data_object,
                    state=state,
                    saved_assets=saved_assets,
                )

        state.unsupported_elements = _unsupported_counts(document.element.body)
        main_text = "\n\n".join(state.text_parts)
        title, title_citations = _document_title(document, state)
        extraction = {
            "document_type": "document",
            "language": None,
            "title": title,
            "title_citations": title_citations,
            "main_text": main_text,
            "main_text_citations": state.text_citations,
            "tables": state.tables,
            "figures": state.figures,
            "formulas": [],
            "formulas_citations": [],
        }
        unsupported_count = sum(state.unsupported_elements.values())
        unsupported_content_count = sum(
            count
            for name, count in state.unsupported_elements.items()
            if name != "drawing_shape"
        )
        return ParsedData(
            object_id=data_object.object_id,
            source_uri=data_object.uri,
            source_format=source_format,
            rows=[
                {
                    "extraction": extraction,
                    "text": main_text,
                    "source_blocks": state.source_blocks,
                    "reading_order": [
                        block["component_id"] for block in state.source_blocks
                    ],
                }
            ],
            text=main_text,
            tables=state.parsed_tables,
            metadata={
                "parser": self.backend_name,
                "page_count": None,
                "logical_page_count": 1,
                "pagination_source": "unavailable",
                "source_block_count": len(state.source_blocks),
                "paragraph_count": len(state.text_parts),
                "body_paragraph_count": len(document.paragraphs),
                "table_count": len(state.tables),
                "drawing_count": sum(
                    1 for _ in _descendants(document.element.body, "drawing")
                ),
                "figure_count": len(state.figures),
                "formula_count": 0,
                "image_count": sum(
                    item.get("status") == "saved" for item in state.image_files
                ),
                "image_files": state.image_files,
                "unsupported_element_count": unsupported_count,
                "unsupported_content_element_count": unsupported_content_count,
                "unsupported_element_types": dict(
                    sorted(state.unsupported_elements.items())
                ),
                "reading_order_source": "docx_native",
                "reading_order_complete": unsupported_content_count == 0,
                "headers_footers_included": False,
            },
        )

    def _parse_paragraph(
        self,
        paragraph: Any,
        *,
        document: Any,
        page_box: list[int] | None,
        data_object: DataObject,
        state: _WordState,
        saved_assets: dict[str, dict[str, Any]],
    ) -> None:
        text = _paragraph_text(paragraph)
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if text:
            block_type = "Heading" if _is_heading(style_name) else "Paragraph"
            component_id = _append_block(
                state,
                block_type=block_type,
                text=text,
                page_box=page_box,
            )
            state.text_parts.append(text)
            state.text_citations.append(component_id)
            if state.first_paragraph is None:
                state.first_paragraph = (text.splitlines()[0], component_id)
            if state.heading_title is None and _is_title_style(style_name):
                state.heading_title = (text.splitlines()[0], component_id)

        self._parse_drawings(
            paragraph._p,
            document=document,
            page_box=page_box,
            data_object=data_object,
            state=state,
            saved_assets=saved_assets,
        )

    def _parse_table(
        self,
        table: Any,
        *,
        document: Any,
        page_box: list[int] | None,
        data_object: DataObject,
        state: _WordState,
        saved_assets: dict[str, dict[str, Any]],
    ) -> None:
        raw_rows = _table_rows(table)
        width = max((len(row) for row in raw_rows), default=0)
        if width and any(cell for row in raw_rows for cell in row):
            raw_rows = [row + [""] * (width - len(row)) for row in raw_rows]
            raw_headers = raw_rows[0]
            headers = normalize_headers(raw_headers, width=width)
            rows = raw_rows[1:]
            content = _markdown_table(raw_headers, rows)
            component_id = _append_block(
                state,
                block_type="Table",
                text=content,
                page_box=page_box,
            )
            table_number = len(state.tables) + 1
            state.tables.append(
                {
                    "caption": "",
                    "caption_citations": [],
                    "content": content,
                    "content_citations": [component_id],
                    "content_format": "markdown",
                    "headers": headers,
                    "rows": rows,
                    "source_ref": component_id,
                }
            )
            state.parsed_tables.append(
                ParsedTable(
                    name=f"Table {table_number}",
                    source_ref=component_id,
                    headers=headers,
                    rows=rows,
                    metadata={"parser": self.backend_name, "logical_page": 0},
                )
            )

        self._parse_drawings(
            table._tbl,
            document=document,
            page_box=page_box,
            data_object=data_object,
            state=state,
            saved_assets=saved_assets,
        )

    def _parse_drawings(
        self,
        element: Any,
        *,
        document: Any,
        page_box: list[int] | None,
        data_object: DataObject,
        state: _WordState,
        saved_assets: dict[str, dict[str, Any]],
    ) -> None:
        for drawing in _descendants(element, "drawing"):
            blip = next(iter(_descendants(drawing, "blip")), None)
            relationship_id = _relationship_id(blip)
            if not relationship_id:
                continue
            image_part = document.part.related_parts.get(relationship_id)
            blob = bytes(getattr(image_part, "blob", b""))
            if not blob:
                continue
            description = _drawing_description(drawing)
            extent = next(iter(_descendants(drawing, "extent")), None)
            bbox = _extent_bbox(extent)
            component_id = _append_block(
                state,
                block_type="Figure",
                text=description,
                page_box=page_box,
                bbox=bbox,
            )
            state.figures.append(
                {
                    "caption": "",
                    "caption_citations": [],
                    "description": description,
                    "description_citations": [component_id],
                    "source_ref": component_id,
                }
            )
            digest = hashlib.sha256(blob).hexdigest()
            extension = _image_extension(image_part)
            file_name = f"{digest}.{extension}"
            saved = saved_assets.get(digest)
            if saved is None:
                saved = self._save_image(data_object.object_id, file_name, blob)
                saved_assets[digest] = saved
            state.image_files.append(
                {
                    **saved,
                    "page": 0,
                    "source_ref": component_id,
                    "sha256": digest,
                    "geometry_available": bbox is not None and page_box is not None,
                }
            )

    def _save_image(self, object_id: str, file_name: str, blob: bytes) -> dict[str, Any]:
        if not self.config.extract_images or not self.config.output_dir:
            return {"name": file_name, "path": None, "status": "not_saved"}
        image_dir = Path(self.config.output_dir) / object_id / "images"
        image_dir.mkdir(parents=True, exist_ok=True)
        output_path = image_dir / file_name
        if not output_path.exists():
            temporary = output_path.with_name(f".{output_path.name}.tmp")
            temporary.write_bytes(blob)
            temporary.replace(output_path)
        return {
            "name": file_name,
            "path": portable_path(output_path, self.config.project_root),
            "status": "saved",
        }


def _append_block(
    state: _WordState,
    *,
    block_type: str,
    text: str,
    page_box: list[int] | None,
    bbox: list[int] | None = None,
) -> str:
    block_index = len(state.source_blocks)
    component_id = f"/page/0/{block_type}/{block_index}"
    block: dict[str, Any] = {
        "component_id": component_id,
        "page": 0,
        "block_index": block_index,
        "type": block_type,
        "text": text,
        "source": AXIOM_NATIVE_BLOCK_SOURCE,
        "parser_source": "docx_native",
    }
    if bbox is not None:
        block["bbox"] = bbox
    if page_box is not None:
        block["page_bbox"] = list(page_box)
    state.source_blocks.append(block)
    return component_id


def _paragraph_text(paragraph: Any) -> str:
    text = _clean_text(paragraph.text)
    if not text:
        return ""
    if _is_numbered_paragraph(paragraph):
        level = _numbering_level(paragraph)
        return f"{'  ' * level}- {text}"
    return text


def _is_numbered_paragraph(paragraph: Any) -> bool:
    properties = getattr(paragraph._p, "pPr", None)
    return properties is not None and getattr(properties, "numPr", None) is not None


def _numbering_level(paragraph: Any) -> int:
    properties = getattr(paragraph._p, "pPr", None)
    num_properties = getattr(properties, "numPr", None)
    level = getattr(num_properties, "ilvl", None)
    try:
        return max(int(level.val), 0) if level is not None else 0
    except (TypeError, ValueError, AttributeError):
        return 0


def _is_heading(style_name: str) -> bool:
    normalized = style_name.strip().casefold()
    return normalized == "title" or normalized.startswith("heading ")


def _is_title_style(style_name: str) -> bool:
    normalized = style_name.strip().casefold()
    return normalized in {"title", "heading 1"}


def _document_title(document: Any, state: _WordState) -> tuple[str | None, list[str]]:
    core_title = _clean_text(getattr(document.core_properties, "title", ""))
    if core_title and core_title.casefold() not in _GENERIC_CORE_TITLES:
        return core_title, []
    if state.heading_title is not None:
        return state.heading_title[0], [state.heading_title[1]]
    if state.first_paragraph is not None:
        return state.first_paragraph[0], [state.first_paragraph[1]]
    return None, []


def _table_rows(table: Any) -> list[list[str]]:
    # Keep the XML elements themselves alive. Storing only ``id(_tc)`` lets
    # Python reuse an id after a short-lived Cell proxy is collected.
    seen_origins: set[Any] = set()
    rows: list[list[str]] = []
    for row in table.rows:
        values: list[str] = []
        for cell in row.cells:
            cell_key = cell._tc
            if cell_key in seen_origins:
                values.append("")
            else:
                seen_origins.add(cell_key)
                values.append(_clean_text(cell.text))
        rows.append(values)
    return rows


def _markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    def render(row: Iterable[str]) -> str:
        return "| " + " | ".join(_markdown_cell(cell) for cell in row) + " |"

    return "\n".join(
        [render(headers), render(["---"] * len(headers)), *(render(row) for row in rows)]
    )


def _markdown_cell(value: Any) -> str:
    return str(value or "").replace("\\", "\\\\").replace("|", "\\|").replace(
        "\n", "<br>"
    )


def _page_box(document: Any) -> list[int] | None:
    if not document.sections:
        return None
    section = document.sections[0]
    width = getattr(section, "page_width", None)
    height = getattr(section, "page_height", None)
    if width is None or height is None:
        return None
    return [0, 0, int(width), int(height)]


def _descendants(element: Any, local_name: str) -> Iterable[Any]:
    if element is None:
        return ()
    return (
        child
        for child in element.iter()
        if str(child.tag).rsplit("}", 1)[-1] == local_name
    )


def _relationship_id(blip: Any) -> str | None:
    if blip is None:
        return None
    for key, value in blip.attrib.items():
        if str(key).rsplit("}", 1)[-1] == "embed":
            return str(value)
    return None


def _drawing_description(drawing: Any) -> str:
    for doc_properties in _descendants(drawing, "docPr"):
        for field in ("descr", "title"):
            value = str(doc_properties.get(field) or "").strip()
            if value:
                return value
    return ""


def _extent_bbox(extent: Any) -> list[int] | None:
    if extent is None:
        return None
    try:
        width = int(extent.get("cx"))
        height = int(extent.get("cy"))
    except (TypeError, ValueError):
        return None
    if width <= 0 or height <= 0:
        return None
    return [0, 0, width, height]


def _image_extension(image_part: Any) -> str:
    part_name = str(getattr(image_part, "partname", ""))
    suffix = Path(part_name).suffix.lower().lstrip(".")
    if suffix == "jpeg":
        return "jpg"
    if re.fullmatch(r"[a-z0-9]+", suffix):
        return suffix
    guessed = mimetypes.guess_extension(str(getattr(image_part, "content_type", "")))
    return (guessed or ".bin").lower().lstrip(".")


def _unsupported_counts(body: Any) -> dict[str, int]:
    counts: dict[str, int] = {}
    for element in body.iter():
        local_name = str(element.tag).rsplit("}", 1)[-1]
        if local_name in _CONTENT_ELEMENTS:
            counts[local_name] = counts.get(local_name, 0) + 1
        if local_name != "drawing":
            continue
        child_names = {
            str(child.tag).rsplit("}", 1)[-1]
            for child in element.iter()
            if isinstance(child.tag, str)
        }
        if "blip" in child_names:
            continue
        if "chart" in child_names:
            kind = "chart"
        elif "relIds" in child_names:
            kind = "diagram"
        elif "txbxContent" in child_names:
            kind = "textbox_drawing"
        else:
            kind = "drawing_shape"
        counts[kind] = counts.get(kind, 0) + 1
    return counts


def _find_soffice(configured_path: str | None) -> Path | None:
    if configured_path:
        path = Path(configured_path).expanduser()
        if not path.is_file():
            raise RuntimeError(
                f"Configured LibreOffice executable does not exist: {configured_path}"
            )
        return path
    for command in ("soffice.com", "soffice", "libreoffice"):
        resolved = shutil.which(command)
        if resolved:
            return Path(resolved)
    for path in (
        Path(r"C:\Program Files\LibreOffice\program\soffice.com"),
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.com"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ):
        if path.is_file():
            return path
    return None


def _converted_docx(output_dir: Path, stem: str) -> Path | None:
    expected = output_dir / f"{stem}.docx"
    if expected.is_file():
        return expected
    stem_key = stem.casefold()
    return next(
        (
            candidate
            for candidate in output_dir.glob("*.docx")
            if candidate.stem.casefold() == stem_key
        ),
        None,
    )


def _clean_text(value: Any) -> str:
    return "\n".join(
        line.rstrip() for line in str(value or "").replace("\r", "").split("\n")
    ).strip()


def _optional_str(value: Any) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    return text or None
